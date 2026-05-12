#!/usr  /bin/env python3
import os
import sys
import json
import base64
import io
import re
import tempfile
import mimetypes
import subprocess
import shutil
import argparse
from pathlib import Path
from datetime import datetime, timezone
from typing import NoReturn, Any, Dict
from memory import (
    load_chat_data,
    save_chat_data,
    get_chat_messages,
    add_memory_chunk,
    update_code_context_from_message,
    build_memory_context,
    should_update_summary,
    make_simple_summary,
)

import requests

CONFIG_PATH = Path.home() / ".config" / "capture-ai" / "config.json"
CACHE_BASE = Path.home() / ".cache" / "capture-ai"
REFS_CACHE_DIR = CACHE_BASE / "refs"
GENERATED_FILES_DIR = CACHE_BASE / "generated_files"
GENERATED_FILES_DIR.mkdir(parents=True, exist_ok=True)
LANG_DIR = Path.home() / "capture-ai" / "language"

_LANG_CACHE = None
_LANG_CACHE_LANG = None

OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"


PROMPT_BLOCK_KEYS = {
    "copyable": "o_Prompt_Copyable",
    "apply": "o_Prompt_Apply",
    "file_create": "o_Prompt_File_Create",
    "web_search": "o_Prompt_Web_Search",
    "structured": "o_Prompt_Structured",
    "code": "o_Prompt_Code",
}

PDF_RULE_KEYS = {
    "pdf_image": "o_Prompt_PDF_Image_Return_Rule",
    "pdf_text_only_to_docx": "o_Prompt_PDF_Text_Only_To_DOCX_Rule",
    "pdf_mixed_separate": "o_Prompt_PDF_Mixed_Separate_Rule",
    "pdf_mixed_text_only": "o_Prompt_PDF_Mixed_Text_Only_Rule",
}


def get_prompt_text(cfg: dict, key: str) -> str:
    text = get_ui_text(cfg, key)
    return "" if text == key else str(text).strip()


def build_global_system_prompt(selected_blocks, cfg: dict) -> str:
    if not isinstance(selected_blocks, list):
        selected_blocks = ["copyable"]

    parts = []

    base = get_prompt_text(cfg, "o_Base_System_Prompt")
    if base:
        parts.append(base)

    for block_name in selected_blocks:
        block_name = str(block_name or "").strip()
        lang_key = PROMPT_BLOCK_KEYS.get(block_name)

        if not lang_key:
            continue

        block_text = get_prompt_text(cfg, lang_key)
        if block_text:
            parts.append(block_text)

    return "\n\n".join(parts).strip()

def load_config():
    try:
        if CONFIG_PATH.exists():
            with open(CONFIG_PATH, "r", encoding="utf-8") as f:
                return json.load(f) or {}
    except Exception as e:
        print("config read error:", e)
    return {}

def _read_ui_language_from_config() -> str:
    try:
        if CONFIG_PATH.exists():
            raw = json.loads(CONFIG_PATH.read_text(encoding="utf-8")) or {}
            return str(raw.get("ui_language", "en")).strip().lower() or "en"
    except Exception:
        pass
    return "en"

def _load_language_map(lang: str) -> dict:
    global _LANG_CACHE, _LANG_CACHE_LANG

    lang = str(lang or "").strip().lower() or "en"

    if _LANG_CACHE is not None and _LANG_CACHE_LANG == lang:
        return _LANG_CACHE

    lang_path = LANG_DIR / f"{lang}.json"
    fallback_path = LANG_DIR / "en.json"

    data = {}

    try:
        if lang_path.exists():
            raw = json.loads(lang_path.read_text(encoding="utf-8")) or {}
            if isinstance(raw, dict):
                data = raw
    except Exception as e:
        print(f"language file read error ({lang_path.name}):", e)

    if not data:
        try:
            if fallback_path.exists():
                raw = json.loads(fallback_path.read_text(encoding="utf-8")) or {}
                if isinstance(raw, dict):
                    data = raw
        except Exception as e:
            print(f"fallback language file read error ({fallback_path.name}):", e)

    if not isinstance(data, dict):
        data = {}

    _LANG_CACHE = data
    _LANG_CACHE_LANG = lang
    return _LANG_CACHE

def get_ui_text(cfg: dict, key: str, **kwargs) -> str:
    lang = str((cfg or {}).get("ui_language", "")).strip().lower()
    if not lang:
        lang = _read_ui_language_from_config()

    lang_map = _load_language_map(lang)

    if not lang_map:
        return key

    text = str(lang_map.get(key, key))

    try:
        return text.format(**kwargs)
    except Exception:
        return text

# ---------------------WEB SEARCH---------------------
def _tavily_search(cfg: dict, query: str, max_results: int = 5) -> str:
    api_key = str(cfg.get("tavily_api_key") or "").strip()
    if not api_key:
        return "[WEB_SEARCH_ERROR]\nTavily API key is missing.\n[/WEB_SEARCH_ERROR]"

    query = str(query or "").strip()
    if not query:
        return "[WEB_SEARCH_ERROR]\nEmpty search query.\n[/WEB_SEARCH_ERROR]"

    try:
        r = requests.post(
            "https://api.tavily.com/search",
            headers={
                "Content-Type": "application/json"
            },
            json={
                "api_key": api_key,
                "query": query,
                "search_depth": "basic",
                "max_results": max_results,
                "include_answer": False,
                "include_raw_content": False
            },
            timeout=45
        )

        if r.status_code != 200:
            try:
                err_json = r.json()
                err_text = (
                    err_json.get("detail")
                    or err_json.get("error")
                    or err_json.get("message")
                    or ""
                )
            except Exception:
                err_text = ""

            if r.status_code in (401, 403):
                msg = get_ui_text(
                    cfg,
                    "o_Web_Search_API_Key_Invalid"
                )

            else:
                msg = get_ui_text(
                    cfg,
                    "o_Web_Search_API_Error",
                    status_code=r.status_code
                )

                if err_text:
                    msg += f"\n{err_text}"

            return (
                "[WEB_SEARCH_ERROR]\n"
                f"{msg}\n"
                "[/WEB_SEARCH_ERROR]"
            )

        data = r.json()
        results = data.get("results") or []

        lines = ["[WEB_SEARCH_RESULTS]", f"query: {query}"]

        for i, item in enumerate(results[:max_results], start=1):
            title = str(item.get("title") or "").strip()
            url = str(item.get("url") or "").strip()
            content = str(item.get("content") or "").strip()

            lines.append(f"\n{i}. {title}")
            lines.append(f"url: {url}")
            lines.append(f"snippet: {content}")

        lines.append("[/WEB_SEARCH_RESULTS]")
        return "\n".join(lines)

    except Exception as e:
        return f"[WEB_SEARCH_ERROR]\n{e}\n[/WEB_SEARCH_ERROR]"

def _extract_web_search_call(text: str) -> str | None:
    raw = str(text or "").strip()

    try:
        data = json.loads(raw)
    except Exception:
        m = re.search(r"\{.*?\"tool\"\s*:\s*\"web_search\".*?\}", raw, re.S)
        if not m:
            return None

        try:
            data = json.loads(m.group(0))
        except Exception:
            return None

    if not isinstance(data, dict):
        return None

    if str(data.get("tool") or "").strip() != "web_search":
        return None

    query = str(data.get("query") or "").strip()
    return query or None




def _safe_read_text(path: Path, max_bytes: int = 250_000) -> str:
    try:
        data = path.read_bytes()
        if len(data) > max_bytes:
            data = data[:max_bytes]
        return data.decode("utf-8", errors="replace")
    except Exception as e:
        return f"[Dosya okunamadı: {e}]"

def _safe_read_pdf(path: Path, max_chars: int = 50000) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        parts = []

        for i, page in enumerate(reader.pages[:20], start=1):
            txt = (page.extract_text() or "").strip()
            if txt:
                parts.append(f"\n--- PDF PAGE {i} ---\n{txt}")

        return "\n".join(parts).strip()[:max_chars]

    except Exception as e:
        print("pdf text read error:", e, file=sys.stderr)
        return ""

def _looks_like_empty_pdf_text(text: str) -> bool:
    t = (text or "").strip()
    return len(t) < 80

def _user_wants_image_edit(text: str) -> bool:
    t = str(text or "").lower()

    image_words = [
        "image", "picture", "photo", "graphic", "diagram", "visual",
        "resim", "görsel", "gorsel", "fotoğraf", "fotograf",
        "şekil", "sekil", "grafik", "diyagram"
    ]

    edit_words = [
        "change", "replace", "modify", "edit", "remove", "add",
        "değiştir", "degistir", "düzenle", "duzenle",
        "kaldır", "kaldir", "ekle", "yenile"
    ]

    return any(w in t for w in image_words) and any(w in t for w in edit_words)


def _user_wants_text_edit(text: str) -> bool:
    t = str(text or "").lower()

    text_words = [
        "text", "writing", "font", "title", "heading", "paragraph",
        "yazı", "yazi", "metin", "font", "başlık", "baslik",
        "paragraf", "içerik", "icerik"
    ]

    edit_words = [
        "shorten", "minimal", "rewrite", "summarize", "increase", "make bigger",
        "kısalt", "kisalt", "minimal", "özetle", "ozetle",
        "büyüt", "buyut", "düzenle", "duzenle"
    ]

    return any(w in t for w in text_words) and any(w in t for w in edit_words)

def _analyze_pdf_kind(path: Path) -> dict:
    """
    PDF türünü ayırır:
    - text_only
    - image_only
    - mixed

    PyMuPDF çalışmazsa pypdf fallback kullanır.
    """
    # 1) Önce PyMuPDF dene
    try:
        import fitz

        doc = fitz.open(str(path))

        text_parts = []
        image_count = 0

        for page in doc:
            txt = (page.get_text("text") or "").strip()
            if txt:
                text_parts.append(txt)

            image_count += len(page.get_images(full=True))

        doc.close()

        text = "\n\n".join(text_parts).strip()
        has_text = len(text) >= 80
        has_images = image_count > 0

        if has_text and has_images:
            kind = "mixed"
        elif has_text:
            kind = "text_only"
        else:
            kind = "image_only"

        return {
            "kind": kind,
            "text": text,
            "has_text": has_text,
            "has_images": has_images,
            "image_count": image_count,
        }

    except Exception as e:
        print("pdf analyze fitz error:", e, file=sys.stderr)

    # 2) Fallback: pypdf ile sadece text var mı bak
    try:
        text = _safe_read_pdf(path)
        has_text = len((text or "").strip()) >= 80

        if has_text:
            # Görsel var mı kesin bilemeyiz, ama en azından image_only'e düşürmeyiz.
            return {
                "kind": "text_only",
                "text": text,
                "has_text": True,
                "has_images": False,
                "image_count": 0,
            }

    except Exception as e:
        print("pdf analyze pypdf fallback error:", e, file=sys.stderr)

    return {
        "kind": "image_only",
        "text": "",
        "has_text": False,
        "has_images": False,
        "image_count": 0,
    }


def _extract_pdf_mixed_layout(path: Path) -> dict:
    """
    Mixed PDF için text ve image bloklarını koordinatlarıyla çıkarır.
    Image bloklarını ayrıca dosya olarak kaydeder.
    """
    try:
        import fitz

        doc = fitz.open(str(path))

        image_dir = GENERATED_FILES_DIR / "pdf_layout_images" / path.stem
        image_dir.mkdir(parents=True, exist_ok=True)

        layout = {
            "source_path": str(path.resolve()),
            "pages": [],
        }

        for page_index, page in enumerate(doc):
            page_dict = page.get_text("dict")
            page_rect = page.rect

            page_info = {
                "page": page_index + 1,
                "width": float(page_rect.width),
                "height": float(page_rect.height),
                "blocks": [],
            }

            for block_index, block in enumerate(page_dict.get("blocks", [])):
                btype = block.get("type")
                bbox = block.get("bbox")

                if not bbox or len(bbox) != 4:
                    continue

                block_id = f"p{page_index + 1}_b{block_index}"

                # Text block
                if btype == 0:
                    text_parts = []

                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            t = str(span.get("text") or "")
                            if t.strip():
                                text_parts.append(t)

                    text = " ".join(text_parts).strip()

                    if text:
                        page_info["blocks"].append({
                            "id": block_id,
                            "type": "text",
                            "bbox": bbox,
                            "text": text,
                        })

                # Image block
                elif btype == 1:
                    image_bytes = block.get("image")
                    ext = str(block.get("ext") or "png").lower().strip() or "png"

                    if image_bytes:
                        img_path = image_dir / f"{block_id}.{ext}"
                        img_path.write_bytes(image_bytes)

                        page_info["blocks"].append({
                            "id": block_id,
                            "type": "image",
                            "bbox": bbox,
                            "path": str(img_path.resolve()),
                        })

            layout["pages"].append(page_info)

        doc.close()
        return layout

    except Exception as e:
        print("mixed pdf layout extract error:", e, file=sys.stderr)
        return {
            "source_path": str(path),
            "pages": [],
        }


def _create_pdf_from_mixed_layout(
    layout: dict,
    text_replacements: dict | None,
    image_replacements: dict | None,
    output_name: str = "edited-mixed-pdf.pdf"
) -> Path | None:
    """
    Mixed PDF'i yeniden kurar:
    - Eski resimleri eski koordinatlarına koyar.
    - AI'den gelen text_replacements ile textleri değiştirir.
    """
    try:
        import fitz

        text_replacements = text_replacements if isinstance(text_replacements, dict) else {}
        image_replacements = image_replacements if isinstance(image_replacements, dict) else {}

        out = _generated_output_path(output_name, "pdf")
        new_doc = fitz.open()

        for page_info in layout.get("pages", []):
            width = float(page_info.get("width") or 595)
            height = float(page_info.get("height") or 842)

            page = new_doc.new_page(width=width, height=height)

            # Önce resimler
            for block in page_info.get("blocks", []):
                if block.get("type") != "image":
                    continue

                bbox = block.get("bbox")
                if not bbox or len(bbox) != 4:
                    continue

                block_id = str(block.get("id") or "")
                img_path = str(image_replacements.get(block_id) or block.get("path") or "")

                if not img_path or not Path(img_path).exists():
                    continue

                try:
                    page.insert_image(fitz.Rect(*bbox), filename=img_path)
                except Exception as e:
                    print("mixed insert image error:", e, file=sys.stderr)

            # Sonra textler
            for block in page_info.get("blocks", []):
                if block.get("type") != "text":
                    continue

                bbox = block.get("bbox")
                if not bbox or len(bbox) != 4:
                    continue

                block_id = str(block.get("id") or "")
                old_text = str(block.get("text") or "")
                new_text = str(text_replacements.get(block_id, old_text) or "")

                try:
                    page.insert_textbox(
                        fitz.Rect(*bbox),
                        new_text,
                        fontsize=11,
                        fontname="helv",
                        color=(0, 0, 0),
                        align=0,
                    )
                except Exception as e:
                    print("mixed insert text error:", e, file=sys.stderr)

        new_doc.save(str(out))
        new_doc.close()

        return out if out.exists() else None

    except Exception as e:
        print("mixed pdf create error:", e, file=sys.stderr)
        return None


def _pdf_to_page_images(path: Path, max_pages: int = 1, zoom: float = 1.5) -> list[Path]:
    out_dir = Path(tempfile.gettempdir()) / "capture-ai-pdf-pages" / path.stem
    out_dir.mkdir(parents=True, exist_ok=True)

    # Eski cache kalıntılarını temizle
    for old in out_dir.glob("page*"):
        try:
            old.unlink()
        except Exception:
            pass

    errors = []

    # 1) PyMuPDF / fitz ile dene
    try:
        import fitz

        doc = fitz.open(str(path))
        out = []

        page_count = min(len(doc), max_pages)

        if page_count <= 0:
            errors.append("fitz: PDF has 0 pages")
        else:
            for i in range(page_count):
                page = doc[i]

                pix = page.get_pixmap(
                    matrix=fitz.Matrix(zoom, zoom),
                    alpha=False
                )

                img_path = out_dir / f"page_{i + 1}.png"
                pix.save(str(img_path))

                if img_path.exists() and img_path.stat().st_size > 0:
                    out.append(img_path)
                else:
                    errors.append(f"fitz: page {i + 1} produced empty PNG")

        doc.close()

        if out:
            return out

    except Exception as e:
        errors.append(f"fitz: {e}")

    # 2) Fallback: poppler / pdftoppm ile dene
    try:
        pdftoppm = shutil.which("pdftoppm")

        if not pdftoppm:
            errors.append("pdftoppm: command not found")
        else:
            prefix = out_dir / "page"

            cmd = [
                pdftoppm,
                "-png",
                "-r", "150",
                "-f", "1",
                "-l", str(max_pages),
                str(path),
                str(prefix),
            ]

            proc = subprocess.run(
                cmd,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.PIPE,
                text=True
            )

            if proc.returncode != 0:
                errors.append(f"pdftoppm: {proc.stderr.strip()}")
            else:
                out = sorted(out_dir.glob("page-*.png"))
                out = [
                    p for p in out
                    if p.exists() and p.stat().st_size > 0
                ]

                if out:
                    return out[:max_pages]

                errors.append("pdftoppm: no PNG output")

    except Exception as e:
        errors.append(f"pdftoppm: {e}")

    print("PDF_CONVERT_FAILED:", " | ".join(errors), file=sys.stderr)
    return []

def _safe_read_docx(path: Path, max_chars: int = 50000) -> str:
    try:
        import docx

        doc = docx.Document(str(path))
        parts = []

        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)

        # tabloları da oku
        for table in doc.tables:
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    txt = (cell.text or "").strip()
                    if txt:
                        cells.append(txt)
                if cells:
                    parts.append(" | ".join(cells))

        out = "\n".join(parts).strip()
        return out[:max_chars] if out else "[DOCX file is empty or text could not be extracted.]"

    except Exception as e:
        return f"[DOCX okunamadı: {e}]"

def _safe_read_xlsx(path: Path, max_rows: int = 200, max_chars: int = 50000) -> str:
    try:
        from openpyxl import load_workbook

        wb = load_workbook(str(path), data_only=True)
        parts = []

        for ws in wb.worksheets:
            parts.append(f"\n--- SHEET: {ws.title} ---")

            row_count = 0
            for row in ws.iter_rows(values_only=True):
                vals = ["" if v is None else str(v) for v in row]
                if any(v.strip() for v in vals):
                    parts.append(" | ".join(vals))
                    row_count += 1

                if row_count >= max_rows:
                    parts.append("[Sheet truncated]")
                    break

        return "\n".join(parts).strip()[:max_chars]

    except Exception as e:
        return f"[XLSX okunamadı: {e}]"

def _is_text_file(path: Path) -> bool:
    ext = path.suffix.lower()
    return ext in {
        ".txt", ".md", ".py", ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf",
        ".sh", ".bash", ".zsh", ".js", ".ts", ".tsx", ".jsx", ".css", ".html", ".xml",
        ".c", ".cpp", ".h", ".hpp", ".rs", ".go", ".java", ".kt", ".cs", ".sql"
    }


def _die(msg: str, code: int = 1) -> NoReturn:
    sys.stderr.write(msg + "\n")
    raise SystemExit(code)


def _load_config() -> Dict[str, Any]:
    if not CONFIG_PATH.exists():
        _die(get_ui_text({}, "o_Config_Not_Found"), 1)
    try:
        with open(CONFIG_PATH, "r", encoding="utf-8") as f:
            return json.load(f) or {}
    except Exception as e:
        _die(get_ui_text({}, "o_Config_Read_Error", error=e), 1)

def _save_config(cfg: dict):
    try:
        CONFIG_PATH.parent.mkdir(parents=True, exist_ok=True)
        CONFIG_PATH.write_text(json.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        # config yazılamasa bile AI çalışsın
        pass


def _parse_selected_indexes(arg: str) -> list[int]:
    out = []
    if not arg:
        return out
    for x in arg.split(","):
        x = x.strip()
        if x.isdigit():
            out.append(int(x))

    # uniq + stable
    seen = set()
    uniq = []
    for i in out:
        if i not in seen:
            seen.add(i)
            uniq.append(i)
    return uniq


def _mime_for(path: Path) -> str:
    mime, _ = mimetypes.guess_type(str(path))
    return mime or "application/octet-stream"


def _data_url_for_image(path: Path) -> str:
    mime = _mime_for(path)
    b64 = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:{mime};base64,{b64}"


def _role_map(chat_role: str) -> str:
    return "assistant" if chat_role == "bot" else "user"


def _safe_chat_key(chat_file: Path) -> str:
    return chat_file.stem.replace("/", "_").replace("\\", "_")


def _utc_stamp() -> str:
    return datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")


def _copy_to_cache(src: Path, dest_dir: Path) -> Path:
    dest_dir.mkdir(parents=True, exist_ok=True)
    base = src.stem.replace(" ", "_")
    ext = src.suffix or ".img"
    dest = dest_dir / f"{base}{ext}"
    if dest.exists():
        dest = dest_dir / f"{base}_{_utc_stamp()}{ext}"
    shutil.copy2(src, dest)
    return dest

def _extract_named_block(text: str, name: str) -> str | None:
    pattern = rf"(?ms)^\s*{re.escape(name)}\s*\n(.*?)\n\s*{re.escape(name)}\s*$"
    m = re.search(pattern, text or "")
    return m.group(1).strip() if m else None


def _remove_named_block(text: str, name: str) -> str:
    pattern = rf"(?ms)\n?\s*{re.escape(name)}\s*\n.*?\n\s*{re.escape(name)}\s*"
    return re.sub(pattern, "", text or "").strip()


def _safe_output_name(name: str, fallback_ext: str) -> str:
    name = str(name or "").strip()

    if not name:
        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        name = f"generated-{stamp}.{fallback_ext}"

    for ch in ["/", "\\", "\n", "\r", "\t"]:
        name = name.replace(ch, "_")

    if "." not in Path(name).name:
        name += f".{fallback_ext}"

    return name


def _generated_output_path(output_name: str, fallback_ext: str) -> Path:
    GENERATED_FILES_DIR.mkdir(parents=True, exist_ok=True)

    safe_name = _safe_output_name(output_name, fallback_ext)
    out = GENERATED_FILES_DIR / safe_name

    if out.exists():
        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        out = GENERATED_FILES_DIR / f"{out.stem}-{stamp}{out.suffix}"

    return out

def _message_content_with_generated_files(msg: dict) -> str:
    text = str(msg.get("content") or "").strip()

    lines = []
    if text:
        lines.append(text)

    images = msg.get("images")
    if isinstance(images, list) and images:
        lines.append("\n[GENERATED_IMAGES]")
        for p in images:
            if p:
                lines.append(f"path: {p}")
        lines.append("[/GENERATED_IMAGES]")

    generated = msg.get("generated_files")
    if isinstance(generated, list) and generated:
        lines.append("\n[GENERATED_FILES]")
        for f in generated:
            if not isinstance(f, dict):
                continue

            path = str(f.get("path") or "").strip()
            name = str(f.get("name") or Path(path).name).strip()

            if path:
                lines.append(f"name: {name}")
                lines.append(f"path: {path}")

        lines.append("[/GENERATED_FILES]")

    return "\n".join(lines).strip()

def _sanitize_plain_document_text(content: str) -> str:
    text = str(content or "")

    # Markdown image: ![alt](url)
    text = re.sub(r"!\[[^\]]*\]\([^)]+\)", "", text)

    # HTML img tag
    text = re.sub(r"<\s*img[^>]*>", "", text, flags=re.IGNORECASE)

    # data:image base64 blokları
    text = re.sub(r"data:image\/[a-zA-Z0-9.+-]+;base64,[A-Za-z0-9+/=\s]+", "", text)

    # Basit HTML tag temizliği
    text = re.sub(r"<[^>]+>", "", text)

    # Çoklu boş satırları sadeleştir
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()

def _create_docx_file(output_name: str, content: str) -> Path | None:
    try:
        import docx
        import docx.shared

        out = _generated_output_path(output_name, "docx")
        doc = docx.Document()

        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = docx.shared.Pt(12)

        section = doc.sections[0]
        section.top_margin = docx.shared.Cm(2)
        section.bottom_margin = docx.shared.Cm(2)
        section.left_margin = docx.shared.Cm(2)
        section.right_margin = docx.shared.Cm(2)

        content = _sanitize_plain_document_text(content)

        for line in str(content or "").splitlines():
            doc.add_paragraph(line)

        doc.save(str(out))
        return out

    except Exception as e:
        print("docx create error:", e, file=sys.stderr)
        return None

def _convert_docx_to_pdf(docx_path: Path) -> Path | None:
    """
    DOCX dosyasını LibreOffice ile PDF'e çevirir.
    """
    try:
        if not docx_path.exists():
            return None

        GENERATED_FILES_DIR.mkdir(parents=True, exist_ok=True)

        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(GENERATED_FILES_DIR),
                str(docx_path),
            ],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )

        converted = GENERATED_FILES_DIR / f"{docx_path.stem}.pdf"

        if converted.exists():
            return converted

        return None

    except Exception as e:
        print("docx to pdf convert error:", e, file=sys.stderr)
        return None

def _norm_filter_text(value: str) -> str:
    import unicodedata

    s = str(value or "").strip().lower()
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))

    tr_map = str.maketrans({
        "ı": "i",
        "ğ": "g",
        "ü": "u",
        "ş": "s",
        "ö": "o",
        "ç": "c",
    })

    return s.translate(tr_map)


def _xlsx_filter_keep_rows(source_path: str, output_name: str, spec: dict) -> Path | None:
    try:
        from openpyxl import load_workbook

        src = Path(source_path).expanduser()
        if not src.exists() or src.suffix.lower() != ".xlsx":
            return None

        out = _generated_output_path(output_name, "xlsx")

        wb = load_workbook(str(src))

        preferred_columns = spec.get("preferred_columns") or []
        patterns = spec.get("patterns_keep_any") or []

        preferred_columns_norm = [
            _norm_filter_text(x) for x in preferred_columns
        ]

        patterns_norm = [
            _norm_filter_text(x) for x in patterns
        ]

        fallback_all = bool(spec.get("fallback_search_all_columns", True))

        for ws in wb.worksheets:
            if ws.max_row < 2:
                continue

            headers = []
            for cell in ws[1]:
                headers.append(_norm_filter_text(cell.value))

            target_indexes = []

            for idx, header in enumerate(headers, start=1):
                if header in preferred_columns_norm:
                    target_indexes.append(idx)

            if not target_indexes and fallback_all:
                target_indexes = list(range(1, ws.max_column + 1))

            rows_to_delete = []

            for row_idx in range(2, ws.max_row + 1):
                values = []

                for col_idx in target_indexes:
                    values.append(ws.cell(row=row_idx, column=col_idx).value)

                joined = _norm_filter_text(" ".join(str(v or "") for v in values))

                keep = any(pat in joined for pat in patterns_norm)

                if not keep:
                    rows_to_delete.append(row_idx)

            for row_idx in reversed(rows_to_delete):
                ws.delete_rows(row_idx, 1)

        wb.save(str(out))
        return out

    except Exception as e:
        print("xlsx filter error:", e, file=sys.stderr)
        return None

def _normalize_xlsx_rows(content):
    rows = content

    # Eğer string olarak Python dict/JSON geldiyse çöz
    if isinstance(rows, str):
        s = rows.strip()

        try:
            rows = json.loads(s)
        except Exception:
            try:
                import ast
                rows = ast.literal_eval(s)
            except Exception:
                return [[rows]]

    # {"sheet": "People", "data": [{...}, {...}]}
    # {"sheet_name": "People", "columns": [...], "rows": [...]}
    # {"sheets": [{"data": [...]}]}
    if isinstance(rows, dict):

        if isinstance(rows.get("sheets"), list):
            first_sheet = rows["sheets"][0] if rows["sheets"] else {}

            if isinstance(first_sheet, dict):
                rows = first_sheet.get("data", [])

        elif isinstance(rows.get("data"), list):
            data_rows = rows.get("data") or []

            # data: [{Name:..., Age:...}, {...}]
            if data_rows and all(isinstance(x, dict) for x in data_rows):
                headers = list(data_rows[0].keys())
                rows = [headers] + [
                    [item.get(h, "") for h in headers]
                    for item in data_rows
                ]
            else:
                rows = data_rows

        elif isinstance(rows.get("rows"), list):
            data_rows = rows.get("rows") or []
            cols = rows.get("columns")

            if isinstance(cols, list):
                rows = [cols] + data_rows
            else:
                rows = data_rows

        else:
            return [[str(rows)]]

    # [{"sheet": "...", "data": [...]}] gibi liste içinde dict geldiyse
    if isinstance(rows, list) and rows and isinstance(rows[0], dict):
        first = rows[0]

        if isinstance(first.get("sheets"), list):
            first_sheet = first["sheets"][0] if first["sheets"] else {}

            if isinstance(first_sheet, dict):
                rows = first_sheet.get("data", [])

        elif isinstance(first.get("data"), list):
            data_rows = first.get("data") or []

            if data_rows and all(isinstance(x, dict) for x in data_rows):
                headers = list(data_rows[0].keys())
                rows = [headers] + [
                    [item.get(h, "") for h in headers]
                    for item in data_rows
                ]
            else:
                rows = data_rows

        elif isinstance(first.get("rows"), list):
            data_rows = first.get("rows") or []
            cols = first.get("columns")

            if isinstance(cols, list):
                rows = [cols] + data_rows
            else:
                rows = data_rows

    if not isinstance(rows, list):
        return [[str(rows)]]

    # [[["Name","Age"],["Ali",20]]] gibi tek satır içine gömülü tablo
    if (
        len(rows) == 1
        and isinstance(rows[0], list)
        and rows[0]
        and all(isinstance(x, list) for x in rows[0])
    ):
        rows = rows[0]

    fixed_rows = []

    for row in rows:
        if isinstance(row, list):
            fixed_rows.append(row)

        elif isinstance(row, dict):
            headers = list(row.keys())

            if not fixed_rows:
                fixed_rows.append(headers)

            fixed_rows.append([row.get(h, "") for h in headers])

        else:
            fixed_rows.append([row])

    return fixed_rows

def _create_xlsx_file(output_name: str, content) -> Path | None:
    try:
        from openpyxl import Workbook

        out = _generated_output_path(output_name, "xlsx")
        wb = Workbook()
        ws = wb.active
        ws.title = "Sheet1"

        rows = _normalize_xlsx_rows(content)

        for row in rows:
            ws.append(row)

        wb.save(str(out))
        return out

    except Exception as e:
        print("xlsx create error:", e, file=sys.stderr)
        return None

def _create_txt_like_file(output_name: str, content: str, ext: str) -> Path | None:
    try:
        out = _generated_output_path(output_name, ext)
        out.write_text(str(content or ""), encoding="utf-8")
        return out

    except Exception as e:
        print(f"{ext} create error:", e, file=sys.stderr)
        return None

def _image_bytes_to_pdf_file(image_bytes: bytes, output_name: str = "edited.pdf") -> Path | None:
    try:
        from PIL import Image

        out = _generated_output_path(output_name, "pdf")

        img = Image.open(io.BytesIO(image_bytes))

        if img.mode in ("RGBA", "LA", "P"):
            img = img.convert("RGB")
        elif img.mode != "RGB":
            img = img.convert("RGB")

        img.save(str(out), "PDF", resolution=100.0)

        return out

    except Exception as e:
        print("image to pdf convert error:", e, file=sys.stderr)
        return None

def _image_paths_to_pdf_file(image_paths: list[Path], output_name: str = "edited.pdf") -> Path | None:
    try:
        from PIL import Image

        valid = [Path(p) for p in image_paths if p and Path(p).exists()]
        if not valid:
            return None

        out = _generated_output_path(output_name, "pdf")
        images = []

        for p in valid:
            img = Image.open(str(p))
            if img.mode in ("RGBA", "LA", "P"):
                img = img.convert("RGB")
            elif img.mode != "RGB":
                img = img.convert("RGB")
            images.append(img)

        if not images:
            return None

        first = images[0]
        rest = images[1:]
        first.save(str(out), "PDF", save_all=True, append_images=rest, resolution=100.0)

        return out if out.exists() else None

    except Exception as e:
        print("images to pdf convert error:", e, file=sys.stderr)
        return None


def _save_result_images_to_files(result_obj: dict) -> list[Path]:
    saved = []

    if not isinstance(result_obj, dict):
        return saved

    def save_raw(raw: bytes, ext: str = "png"):
        out = _generated_output_path(f"edited-image.{ext}", ext)
        out.write_bytes(raw)
        if out.exists():
            saved.append(out)

    # URL images
    for key in ("image", "url"):
        val = result_obj.get(key)
        if isinstance(val, str) and val.strip():
            s = val.strip()
            if s.startswith("data:image/") and "," in s:
                try:
                    head, b64 = s.split(",", 1)
                    ext = head.split("/")[1].split(";")[0] or "png"
                    save_raw(base64.b64decode(b64), ext)
                except Exception:
                    pass
            elif s.startswith(("http://", "https://")):
                try:
                    r = requests.get(s, timeout=120)
                    if r.status_code == 200:
                        save_raw(r.content, "png")
                except Exception:
                    pass

    imgs = result_obj.get("images")
    if isinstance(imgs, list):
        for item in imgs:
            if isinstance(item, str):
                s = item.strip()
                if s.startswith("data:image/") and "," in s:
                    try:
                        head, b64 = s.split(",", 1)
                        ext = head.split("/")[1].split(";")[0] or "png"
                        save_raw(base64.b64decode(b64), ext)
                    except Exception:
                        pass
                elif s.startswith(("http://", "https://")):
                    try:
                        r = requests.get(s, timeout=120)
                        if r.status_code == 200:
                            save_raw(r.content, "png")
                    except Exception:
                        pass

            elif isinstance(item, dict):
                u = item.get("url") or item.get("image_url")
                if isinstance(u, dict):
                    u = u.get("url")
                if isinstance(u, str) and u.strip():
                    try:
                        r = requests.get(u.strip(), timeout=120)
                        if r.status_code == 200:
                            save_raw(r.content, "png")
                    except Exception:
                        pass

    # base64 images
    for key in ("image_base64",):
        b64 = result_obj.get(key)
        if isinstance(b64, str) and b64.strip():
            try:
                save_raw(base64.b64decode(b64.split(",", 1)[-1]), "png")
            except Exception:
                pass

    many_b64 = result_obj.get("images_base64")
    if isinstance(many_b64, list):
        for b64 in many_b64:
            if isinstance(b64, str) and b64.strip():
                try:
                    save_raw(base64.b64decode(b64.split(",", 1)[-1]), "png")
                except Exception:
                    pass

    return saved


def _overlay_images_on_pdf_blocks(source_pdf: Path, layout: dict, replacement_images: list[Path], output_name: str = "edited-pdf-image.pdf") -> Path | None:
    try:
        import fitz

        source_pdf = Path(source_pdf)
        if not source_pdf.exists():
            return None

        doc = fitz.open(str(source_pdf))
        out = _generated_output_path(output_name, "pdf")

        image_blocks = []
        for page in layout.get("pages", []):
            page_no = int(page.get("page") or 1)
            for block in page.get("blocks", []):
                if block.get("type") == "image":
                    image_blocks.append((page_no, block))

        if not image_blocks or not replacement_images:
            doc.close()
            return None

        for idx, img_path in enumerate(replacement_images):
            if idx >= len(image_blocks):
                break

            page_no, block = image_blocks[idx]
            bbox = block.get("bbox")
            if not bbox or len(bbox) != 4:
                continue

            page = doc[page_no - 1]
            page.insert_image(
                fitz.Rect(*bbox),
                filename=str(img_path),
                overlay=True
            )

        doc.save(str(out))
        doc.close()

        return out if out.exists() else None

    except Exception as e:
        print("overlay pdf image error:", e, file=sys.stderr)
        return None

def _save_ai_returned_file(name: str, data: bytes) -> Path | None:
    try:
        ext = Path(name or "").suffix.lower().lstrip(".") or "bin"
        out = _generated_output_path(name, ext)
        out.write_bytes(data)
        return out
    except Exception as e:
        print("ai returned file save error:", e, file=sys.stderr)
        return None


def _apply_ai_returned_files_from_result(result_obj: dict) -> list[dict]:
    generated = []

    if not isinstance(result_obj, dict):
        return generated

    # 1) Base64 dosyalar
    files_b64 = result_obj.get("generated_files_base64")
    if isinstance(files_b64, list):
        for item in files_b64:
            if not isinstance(item, dict):
                continue

            name = str(item.get("name") or item.get("output_name") or "generated.pdf").strip()
            b64 = str(item.get("base64") or item.get("data") or "").strip()

            if b64.startswith("data:") and "," in b64:
                b64 = b64.split(",", 1)[1].strip()

            if not b64:
                continue

            try:
                raw = base64.b64decode(b64)
            except Exception:
                continue

            out = _save_ai_returned_file(name, raw)

            if out and out.exists():
                generated.append({
                    "path": str(out.resolve()),
                    "name": out.name
                })

    # 2) URL dosyalar
    files_url = result_obj.get("generated_files_url")
    if isinstance(files_url, list):
        for item in files_url:
            if not isinstance(item, dict):
                continue

            name = str(item.get("name") or item.get("output_name") or "generated.pdf").strip()
            url = str(item.get("url") or "").strip()

            if not url.startswith(("http://", "https://")):
                continue

            try:
                r = requests.get(url, timeout=120)
                if r.status_code != 200:
                    continue

                out = _save_ai_returned_file(name, r.content)

                if out and out.exists():
                    generated.append({
                        "path": str(out.resolve()),
                        "name": out.name
                    })

            except Exception as e:
                print("ai returned file url download error:", e, file=sys.stderr)

    return generated

def _apply_file_create_from_reply(reply_text: str) -> list[dict]:
    block = _extract_named_block(reply_text, "file_create")
    if not block:
        return []

    try:
        ops = json.loads(block)
    except Exception as e:
        print("file_create json parse error:", e, file=sys.stderr)
        return []

    if isinstance(ops, dict):
        ops = [ops]

    if not isinstance(ops, list):
        return []

    generated = []

    for op in ops:
        if not isinstance(op, dict):
            continue

        fmt = str(op.get("format") or "").strip().lower().lstrip(".")
        output_name = str(op.get("output_name") or "").strip()
        content = op.get("content", "")

        out = None

        if fmt == "docx":
            out = _create_docx_file(output_name, str(content or ""))

            if out and out.exists():
                generated.append({
                    "path": str(out.resolve()),
                    "name": out.name
                })

                pdf_out = _convert_docx_to_pdf(out)
                if pdf_out and pdf_out.exists():
                    generated.append({
                        "path": str(pdf_out.resolve()),
                        "name": pdf_out.name
                    })

                continue

        elif fmt == "xlsx":
            source_path = str(op.get("source_path") or "").strip()

            if isinstance(content, list) and content and isinstance(content[0], dict):
                first_spec = content[0]
                action = str(first_spec.get("action") or "").strip()

                if action == "filter_keep_rows":
                    out = _xlsx_filter_keep_rows(source_path, output_name, first_spec)
                else:
                    out = _create_xlsx_file(output_name, content)
            else:
                out = _create_xlsx_file(output_name, content)

        elif fmt in ("txt", "md"):
            out = _create_txt_like_file(output_name, str(content or ""), fmt)

        if out and out.exists():
            generated.append({
                "path": str(out.resolve()),
                "name": out.name
            })

    return generated

def _build_blocks_and_cache_info(msg: dict, cache_images_dir: Path | None, pdf_mode: str = "auto"):
    blocks = []
    cached_image = None

    # 1) Text
    txt = msg.get("content") or ""
    if txt:
        blocks.append({"type": "text", "text": txt})

    # 2) Images (legacy "image" + new "images")
    imgs = []
    one = msg.get("image")
    if one:
        imgs.append(one)

    many = msg.get("images")
    if isinstance(many, list):
        imgs.extend([str(x) for x in many if x])

    # uniq
    seen = set()
    uniq_imgs = []
    for p in imgs:
        if p not in seen:
            seen.add(p)
            uniq_imgs.append(p)

    for img_raw in uniq_imgs:
        img_path = Path(img_raw)
        if not img_path.exists():
            continue

        if cache_images_dir is not None:
            try:
                cached_image = _copy_to_cache(img_path, cache_images_dir)
            except Exception:
                cached_image = cached_image  # önceki varsa kalsın

        try:
            blocks.append({
                "type": "image_url",
                "image_url": {"url": _data_url_for_image(img_path)}
            })
        except Exception:
            pass

    # 3) Files (new: "files": [{"path","name","edit"}])
    files = msg.get("files") or []
    if isinstance(files, list):
        for f in files:
            if not isinstance(f, dict):
                continue

            p_raw = str(f.get("path") or "").strip()
            if not p_raw:
                continue
            p = Path(p_raw)
            if not p.exists() or not p.is_file():
                continue

            editable = bool(f.get("edit"))
            name = (f.get("name") or p.name or "").strip()
            mime = _mime_for(p)

            file_text = None
            file_type = "unsupported"
            handled_file = False

            if _is_text_file(p):
                file_text = _safe_read_text(p)
                file_type = "text"

            elif p.suffix.lower() == ".pdf":

                # IMAGE MODE override — HER ZAMAN image gönder
                if pdf_mode == "image":
                    blocks.append({
                        "type": "text",
                        "text": (
                            "\n[FILE]\n"
                            f"name: {name}\n"
                            f"mime: {mime}\n"
                            f"editable: {'true' if editable else 'false'}\n"
                            "type: pdf_visual_pages\n"
                            "[/FILE]\n"
                        )
                    })

                    page_images = _pdf_to_page_images(p, max_pages=12, zoom=1.5)

                    if not page_images:
                        raise RuntimeError(
                            get_ui_text(load_config(), "o_PDF_Convert_Failed")
                        )

                    for img_path in page_images:
                        blocks.append({
                            "type": "image_url",
                            "image_url": {
                                "url": _data_url_for_image(img_path)
                            }
                        })

                    handled_file = True
                    continue

                analysis = _analyze_pdf_kind(p)
                pdf_kind = analysis.get("kind")
                pdf_text = analysis.get("text") or ""

                # 1) Sadece yazı PDF
                if pdf_kind == "text_only":
                    file_text = pdf_text
                    file_type = "pdf_text_only"

                    if not str(file_text or "").strip():
                        blocks.append({
                            "type": "text",
                            "text": (
                                get_ui_text(load_config(), "o_PDF_Text_Extract_Failed")
                            )
                        })
                        handled_file = True
                        continue

                # 2) Hem yazı hem resim PDF
                elif pdf_kind == "mixed":
                    layout = _extract_pdf_mixed_layout(p)

                    if pdf_mode == "text":
                        text_blocks = []

                        for page in layout.get("pages", []):
                            for block in page.get("blocks", []):
                                if block.get("type") == "text":
                                    text_blocks.append({
                                        "id": block.get("id"),
                                        "page": page.get("page"),
                                        "text": block.get("text"),
                                    })

                        blocks.append({
                            "type": "text",
                            "text": (
                                "\n[FILE]\n"
                                f"name: {name}\n"
                                f"source_path_for_app_only: {str(p)}\n"
                                f"mime: {mime}\n"
                                f"editable: {'true' if editable else 'false'}\n"
                                "type: pdf_mixed_text_only\n"
                                "note: Images are preserved by the app. Do not regenerate or describe images.\n"
                                "[CONTENT]\n"
                                f"{json.dumps(text_blocks, ensure_ascii=False)}\n"
                                "[/CONTENT]\n"
                                "[/FILE]\n"
                            )
                        })
                        handled_file = True

                    elif pdf_mode == "image":
                        image_blocks = []

                        for page in layout.get("pages", []):
                            for block in page.get("blocks", []):
                                if block.get("type") == "image":
                                    image_blocks.append(block)
                                    img_path = Path(str(block.get("path") or ""))

                                    if img_path.exists():
                                        blocks.append({
                                            "type": "image_url",
                                            "image_url": {
                                                "url": _data_url_for_image(img_path)
                                            }
                                        })

                        if not image_blocks:
                            blocks.append({
                                "type": "text",
                                "text": get_ui_text(load_config(), "o_PDF_No_Images_Found")
                            })

                        handled_file = True

                    else:
                        blocks.append({
                            "type": "text",
                            "text": (
                                "\n[FILE]\n"
                                f"name: {name}\n"
                                f"path: {str(p)}\n"
                                f"mime: {mime}\n"
                                f"editable: {'true' if editable else 'false'}\n"
                                "type: pdf_mixed_separated_layout\n"
                                "[CONTENT]\n"
                                f"{json.dumps(layout, ensure_ascii=False)}\n"
                                "[/CONTENT]\n"
                                "[/FILE]\n"
                            )
                        })
                        handled_file = True

                # 3) Yazı yok / scan / image PDF
                else:
                    blocks.append({
                        "type": "text",
                        "text": (
                            "\n[FILE]\n"
                            f"name: {name}\n"
                            f"path: {str(p)}\n"
                            f"mime: {mime}\n"
                            f"editable: {'true' if editable else 'false'}\n"
                            "type: pdf_visual_pages\n"
                            "note: PDF text could not be extracted. Pages are sent as images.\n"
                            "[/FILE]\n"
                        )
                    })


                    page_images = _pdf_to_page_images(p, max_pages=12, zoom=1.5)
                    handled_file = True

                    if not page_images:
                        raise RuntimeError(
                            get_ui_text(load_config(), "o_PDF_Convert_Failed")
                        )

                    for img_path in page_images:
                        blocks.append({
                            "type": "image_url",
                            "image_url": {
                                "url": _data_url_for_image(img_path)
                            }
                        })


            elif p.suffix.lower() == ".docx":
                file_text = _safe_read_docx(p)
                file_type = "docx_text"

            elif p.suffix.lower() == ".xlsx":
                file_text = _safe_read_xlsx(p)
                file_type = "xlsx_text"

            if file_text is not None:
                blocks.append({
                    "type": "text",
                    "text": (
                        "\n[FILE]\n"
                        f"name: {name}\n"
                        f"source_path_for_app_only: {str(p)}\n"
                        f"mime: {mime}\n"
                        f"editable: {'true' if editable else 'false'}\n"
                        f"type: {file_type}\n"
                        "IMPORTANT: The file content is already extracted below. Do not access the path.\n"
                        "[CONTENT]\n"
                        f"{file_text}\n"
                        "[/CONTENT]\n"
                        "[/FILE]\n"
                    )
                })
            elif not handled_file:
                blocks.append({
                    "type": "text",
                    "text": (
                        "\n[FILE]\n"
                        f"name: {name}\n"
                        f"path: {str(p)}\n"
                        f"mime: {mime}\n"
                        f"editable: {'true' if editable else 'false'}\n"
                        "type: unsupported\n"
                        "note: Bu dosya türü şu an metin olarak okunamıyor.\n"
                        "[/FILE]\n"
                    )
                })

    if not blocks:
        blocks = [{"type": "text", "text": ""}]

    return blocks, (str(cached_image) if cached_image else None)

def fix_mojibake(s: str) -> str:
    if not s:
        return s
    # Mojibake tipik işaretleri
    if ("Ã" not in s) and ("Å" not in s) and ("â" not in s):
        return s
    try:
        # yanlışlıkla latin-1/cp1252 gibi decode edilmiş UTF-8'i düzelt
        return s.encode("latin-1").decode("utf-8")
    except Exception:
        return s

def _message_to_plain_text(content) -> str:
    if isinstance(content, str):
        return content.strip()

    if isinstance(content, list):
        parts = []
        for item in content:
            if not isinstance(item, dict):
                continue
            if item.get("type") == "text":
                txt = item.get("text")
                if isinstance(txt, str) and txt.strip():
                    parts.append(txt.strip())
        return "\n".join(parts).strip()

    return ""

def lang_code_to_name(code: str) -> str:
    m = {
        "tr": "Turkish",
        "en": "English",
        "de": "German",
        "fr": "French",
        "es": "Spanish",
    }
    return m.get(code, code)

def get_ui_language():
    try:
        if CONFIG_PATH.exists():
            with open(CONFIG_PATH, "r", encoding="utf-8") as f:
                cfg = json.load(f) or {}
                return str(cfg.get("ui_language", "en")).strip().lower()
    except:
        pass
    return "en"


def _extract_images_from_message(msg: dict) -> dict:
    """
    OpenRouter / farklı provider cevaplarından image url veya base64 çıkar.
    UI'nin finalize_ai_response() fonksiyonunun anlayacağı formatı döndürür.
    """
    out = {
        "type": "text",
        "content": "",
    }

    if not isinstance(msg, dict):
        return out

    content = msg.get("content")
    text_parts = []
    image_urls = []
    image_b64 = []

    # 1) content string ise
    if isinstance(content, str):
        if content.strip():
            text_parts.append(content.strip())

    # 2) content liste ise
    elif isinstance(content, list):
        for item in content:
            if not isinstance(item, dict):
                continue

            itype = str(item.get("type") or "").strip().lower()

            if itype == "text":
                txt = item.get("text")
                if isinstance(txt, str) and txt.strip():
                    text_parts.append(txt.strip())

            elif itype in ("image_url", "output_image"):
                iu = item.get("image_url")
                if isinstance(iu, str) and iu.strip():
                    image_urls.append(iu.strip())
                elif isinstance(iu, dict):
                    u = iu.get("url")
                    if isinstance(u, str) and u.strip():
                        image_urls.append(u.strip())

                b64 = item.get("b64_json")
                if isinstance(b64, str) and b64.strip():
                    image_b64.append(b64.strip())

                data = item.get("image")
                if isinstance(data, str) and data.strip():
                    image_b64.append(data.strip())

    # 3) message level fallback alanlar
    for key in ("image", "url"):
        val = msg.get(key)
        if isinstance(val, str) and val.strip():
            image_urls.append(val.strip())

    one_b64 = msg.get("image_base64")
    if isinstance(one_b64, str) and one_b64.strip():
        image_b64.append(one_b64.strip())

    imgs = msg.get("images")
    if isinstance(imgs, list):
        for item in imgs:
            if isinstance(item, str) and item.strip():
                image_urls.append(item.strip())
            elif isinstance(item, dict):
                u = item.get("url")
                if isinstance(u, str) and u.strip():
                    image_urls.append(u.strip())

                iu = item.get("image_url")
                if isinstance(iu, str) and iu.strip():
                    image_urls.append(iu.strip())
                elif isinstance(iu, dict):
                    uu = iu.get("url")
                    if isinstance(uu, str) and uu.strip():
                        image_urls.append(uu.strip())

                b64 = item.get("b64_json")
                if isinstance(b64, str) and b64.strip():
                    image_b64.append(b64.strip())

    # uniq
    def _uniq(seq):
        seen = set()
        out2 = []
        for x in seq:
            if x not in seen:
                seen.add(x)
                out2.append(x)
        return out2

    image_urls = _uniq(image_urls)
    image_b64 = _uniq(image_b64)

    out["content"] = "\n".join(text_parts).strip()

    if image_urls:
        out["type"] = "image"
        out["images"] = image_urls

    if image_b64:
        out["type"] = "image"
        out["images_base64"] = image_b64

    if image_urls and image_b64:
        out["type"] = "image"

    return out

def _extract_images_from_anywhere(obj) -> dict:
    """
    OpenRouter / provider response içinde image url veya base64 farklı yerlerde dönebilir.
    Bu fallback tüm response JSON içinde güvenli şekilde image benzeri alanları arar.
    """
    out = {
        "type": "text",
        "content": "",
    }

    image_urls = []
    image_b64 = []

    image_key_names = {
        "image",
        "url",
        "image_url",
        "output_image",
        "b64_json",
        "image_base64",
        "base64",
        "data",
    }

    def add_url(s: str):
        s = str(s or "").strip()
        if not s:
            return

        if s.startswith("data:image/"):
            image_urls.append(s)
            return

        low = s.lower()
        if low.startswith("http://") or low.startswith("https://"):
            if any(x in low for x in [".png", ".jpg", ".jpeg", ".webp", ".gif"]) or "image" in low:
                image_urls.append(s)

    def add_b64(s: str):
        s = str(s or "").strip()
        if not s:
            return

        if s.startswith("data:image/"):
            image_urls.append(s)
            return

        # çok kısa stringleri base64 sanma
        if len(s) < 500:
            return

        # basit base64 karakter kontrolü
        if re.fullmatch(r"[A-Za-z0-9+/=\s]+", s):
            image_b64.append(s)

    def walk(x, parent_key=""):
        if isinstance(x, dict):
            for k, v in x.items():
                key = str(k or "").strip().lower()

                if isinstance(v, str):
                    if key in image_key_names or "image" in key or "b64" in key:
                        if key in ("b64_json", "image_base64", "base64") or "b64" in key:
                            add_b64(v)
                        else:
                            add_url(v)
                            add_b64(v)
                    else:
                        # data:image/... bazen key alakasız olabilir
                        if v.strip().startswith("data:image/"):
                            add_url(v)

                else:
                    walk(v, key)

        elif isinstance(x, list):
            for item in x:
                walk(item, parent_key)

        elif isinstance(x, str):
            if x.strip().startswith("data:image/"):
                add_url(x)

    walk(obj)

    def uniq(seq):
        seen = set()
        result = []
        for item in seq:
            key = item[:200] + str(len(item))
            if key in seen:
                continue
            seen.add(key)
            result.append(item)
        return result

    image_urls = uniq(image_urls)
    image_b64 = uniq(image_b64)

    if image_urls:
        out["type"] = "image"
        out["images"] = image_urls

    if image_b64:
        out["type"] = "image"
        out["images_base64"] = image_b64

    return out

def _split_local_model(model_name: str) -> tuple[str, str]:
    raw = str(model_name or "").strip()
    if not raw or "/" not in raw:
        cfg = load_config()
        _die(get_ui_text(cfg, "o_Invalid_Local_Model_Format", model=raw), 1)

    provider, actual_model = raw.split("/", 1)
    provider = provider.strip().lower()
    actual_model = actual_model.strip()

    if not provider or not actual_model:
        cfg = load_config()
        _die(get_ui_text(cfg, "o_Invalid_Local_Model_Format", model=raw), 1)

    return provider, actual_model

def _is_image_generation_model(model_name: str) -> bool:
    m = str(model_name or "").strip().lower()
    return any(x in m for x in [
        "image",
        "gpt-image",
        "imagen",
        "flux",
        "recraft",
        "stable-diffusion"
    ])

def _normalize_model_entry(item) -> dict | None:
    if isinstance(item, dict):
        mid = str(item.get("id") or "").strip()
        if not mid:
            return None
        return {
            "id": mid,
            "local": bool(item.get("local", False))
        }

    if isinstance(item, str):
        mid = item.strip()
        if not mid:
            return None
        return {
            "id": mid,
            "local": False
        }

    return None


def _get_models_list(cfg: dict) -> list[dict]:
    raw = cfg.get("ai_models", [])
    if not isinstance(raw, list):
        raw = []

    out = []
    seen = set()

    for item in raw:
        norm = _normalize_model_entry(item)
        if not norm:
            continue

        mid = norm["id"]
        if mid in seen:
            continue

        seen.add(mid)
        out.append(norm)

    return out


def _get_chat_model_entry(cfg: dict, chat_file: Path) -> dict:
    models = _get_models_list(cfg)
    if not models:
        _die("ai_models in config is empty. At least one model must be defined.", 1)

    model_ids = {m["id"] for m in models}

    chat_models = cfg.get("chat_models", {})
    if not isinstance(chat_models, dict):
        chat_models = {}

    chat_key = chat_file.name
    selected = chat_models.get(chat_key)

    chosen_id = ""
    chosen_local = False

    if isinstance(selected, dict):
        chosen_id = str(selected.get("id") or "").strip()
        chosen_local = bool(selected.get("local", False))
    elif isinstance(selected, str):
        chosen_id = selected.strip()
        chosen_local = False

    if not chosen_id or chosen_id not in model_ids:
        fallback = models[0]
        chosen_id = fallback["id"]
        chosen_local = bool(fallback.get("local", False))

        chat_models[chat_key] = {
            "id": chosen_id,
            "local": chosen_local
        }
        cfg["chat_models"] = chat_models
        _save_config(cfg)
        return {
            "id": chosen_id,
            "local": chosen_local
        }

    # ai_models içindeki gerçek local değerini koru
    for m in models:
        if m["id"] == chosen_id:
            chosen_local = bool(m.get("local", False))
            break

    fixed_entry = {
        "id": chosen_id,
        "local": chosen_local
    }

    if selected != fixed_entry:
        chat_models[chat_key] = fixed_entry
        cfg["chat_models"] = chat_models
        _save_config(cfg)

    return fixed_entry

def _pick_model(cfg: dict, chat_file: Path, cli_model: str | None) -> dict:
    if cli_model:
        return {
            "id": str(cli_model).strip(),
            "local": False
        }

    return _get_chat_model_entry(cfg, chat_file)

def _coerce_number(val: str):
    s = str(val or "").strip()
    if not s:
        return None
    try:
        if "." in s:
            return float(s)
        return int(s)
    except Exception:
        return None


def _build_local_options(provider_cfg: dict) -> dict:
    opts = {}

    num = _coerce_number(provider_cfg.get("temperature", ""))
    if num is not None:
        opts["temperature"] = num

    num = _coerce_number(provider_cfg.get("top_p", ""))
    if num is not None:
        opts["top_p"] = num

    num = _coerce_number(provider_cfg.get("top_k", ""))
    if num is not None:
        opts["top_k"] = num

    num = _coerce_number(provider_cfg.get("repeat_penalty", ""))
    if num is not None:
        opts["repeat_penalty"] = num

    num = _coerce_number(provider_cfg.get("num_ctx", ""))
    if num is not None:
        opts["num_ctx"] = num

    num = _coerce_number(provider_cfg.get("num_predict", ""))
    if num is not None:
        opts["num_predict"] = num

    keep_alive = str(provider_cfg.get("keep_alive") or "").strip()
    if keep_alive:
        opts["keep_alive"] = keep_alive

    return opts


def _flatten_messages_for_local(final_messages: list[dict]) -> str:
    parts = []

    for msg in final_messages:
        if not isinstance(msg, dict):
            continue

        role = str(msg.get("role") or "").strip().upper()
        content = msg.get("content")

        if isinstance(content, str):
            txt = content.strip()
            if txt:
                parts.append(f"[{role}]\n{txt}")
            continue

        if isinstance(content, list):
            subparts = []
            for item in content:
                if not isinstance(item, dict):
                    continue

                itype = str(item.get("type") or "").strip().lower()

                if itype == "text":
                    txt = str(item.get("text") or "").strip()
                    if txt:
                        subparts.append(txt)

                elif itype == "image_url":
                    iu = item.get("image_url")
                    if isinstance(iu, dict):
                        u = str(iu.get("url") or "").strip()
                    else:
                        u = str(iu or "").strip()

                    if u:
                        subparts.append(f"[IMAGE]\n{u}")

            joined = "\n".join(x for x in subparts if x.strip()).strip()
            if joined:
                parts.append(f"[{role}]\n{joined}")

    return "\n\n".join(parts).strip()


def _call_ollama(provider_cfg: dict, actual_model: str, final_messages: list[dict]) -> str:
    base_url = str(provider_cfg.get("base_url") or "").strip() or "http://127.0.0.1:11434"
    url = base_url.rstrip("/") + "/api/generate"

    prompt = _flatten_messages_for_local(final_messages)

    payload = {
        "model": actual_model,
        "prompt": prompt,
        "stream": False,
    }

    system_prompt = str(provider_cfg.get("system_prompt") or "").strip()
    if system_prompt:
        payload["system"] = system_prompt

    options = _build_local_options(provider_cfg)
    if options:
        payload["options"] = options

    r = requests.post(
        url,
        json=payload,
        timeout=300
    )

    if r.status_code != 200:
        cfg = load_config()
        _die(get_ui_text(cfg, "o_Ollama_API_Error", status_code=r.status_code, details=r.text[:1000]), 1)

    try:
        data = r.json()
    except Exception as e:
        cfg = load_config()
        _die(get_ui_text(cfg, "o_Ollama_JSON_Parse_Error", error=e), 1)

    text = str(data.get("response") or "").strip()
    if not text:
        cfg = load_config()
        _die(get_ui_text(cfg, "o_Ollama_Empty_Response"), 1)

    result_obj = {
        "type": "text",
        "content": fix_mojibake(text)
    }

    prompt_eval = data.get("prompt_eval_count")
    eval_count = data.get("eval_count")

    if prompt_eval is not None or eval_count is not None:
        p = int(prompt_eval or 0)
        c = int(eval_count or 0)
        result_obj["usage"] = {
            "prompt_tokens": p,
            "completion_tokens": c,
            "total_tokens": p + c
        }

    return json.dumps(result_obj, ensure_ascii=False)

def main():
    try:
        if hasattr(sys.stdout, "reconfigure"):
            sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        if hasattr(sys.stderr, "reconfigure"):
            sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

    # Ekstra sağlamlaştırma (bazı ortamlarda fark yaratır)
    os.environ["PYTHONIOENCODING"] = "utf-8"

    # Args
    parser = argparse.ArgumentParser(add_help=False)
    parser.add_argument("chat_file", nargs="?")
    parser.add_argument("selected_indexes", nargs="?")
    parser.add_argument("--model", dest="model", default=None)
    parser.add_argument("--context-mode", dest="context_mode", default="rag")
    parser.add_argument("--approved-web-search-query", dest="approved_web_search_query", default=None)

    if len(sys.argv) < 2:
        sys.exit(1)

    args, _unknown = parser.parse_known_args()
    
    context_mode = str(args.context_mode or "rag").strip().lower()
    if context_mode not in ("rag", "direct"):
        context_mode = "rag"

    if not args.chat_file:
        sys.exit(1)

    chat_file = Path(args.chat_file)
    if not chat_file.exists():
        sys.exit(1)

    selected_indexes = []
    if args.selected_indexes:
        try:
            selected_indexes = _parse_selected_indexes(args.selected_indexes)
        except Exception:
            selected_indexes = []

    # Config
    cfg = _load_config()
    response_style = (cfg.get("response_style") or "").strip()

    # model seçimi (chat’e göre)
    model_entry = _pick_model(cfg, chat_file, args.model)
    model_name = str(model_entry["id"]).strip()
    is_local_model = bool(model_entry.get("local", False))

    local_providers = cfg.get("local_providers", {})
    if not isinstance(local_providers, dict):
        local_providers = {}

    API_KEY = ""
    if not is_local_model:
        API_KEY = (cfg.get("open_router_key") or "").strip()
        if not API_KEY:
            _die(get_ui_text(cfg, "o_OpenRouter_Key_Missing"), 1)

    # read rag settings
    rag_settings = cfg.get("rag_settings", {})
    if not isinstance(rag_settings, dict):
        rag_settings = {}

    def rag_bool(key, default):
        return bool(rag_settings.get(key, default))

    def rag_int(key, default, min_v=0, max_v=100_000):
        try:
            val = int(rag_settings.get(key, default))
        except Exception:
            val = default
        return max(min_v, min(max_v, val))

    ask_for_web_search = bool(cfg.get("ask_for_web_search", True))
    approved_web_search_query = str(args.approved_web_search_query or "").strip()

    recent_message_count = rag_int("recent_message_count", 10, 0, 50)
    retrieved_chunk_count = rag_int("retrieved_chunk_count", 5, 0, 50)
    summary_update_every = rag_int("summary_update_every", 20, 1, 200)
    memory_chunk_max_chars = rag_int("memory_chunk_max_chars", 1200, 100, 100_000)
    summary_max_chars = rag_int("summary_max_chars", 6000, 100, 100_000)
    code_context_max_chars = rag_int("code_context_max_chars", 8000, 100, 100_000)

    use_summary = rag_bool("use_summary", True)
    use_recent_messages = rag_bool("use_recent_messages", True)
    use_retrieval = rag_bool("use_retrieval", True)
    use_code_context = rag_bool("use_code_context", True)
    include_recent_attachments = rag_bool("include_recent_attachments", False)

    # Read chat
    try:
        chat_data = load_chat_data(chat_file)
        messages = get_chat_messages(chat_data)
    except Exception as e:
        _die(get_ui_text(cfg, "o_Chat_Read_Error", error=e), 1)

    if not messages:
        sys.exit(0)


    if messages and isinstance(messages[-1], dict) and messages[-1].get("streaming") and messages[-1].get("role") in ("bot", "assistant"):
        messages = messages[:-1]  # placeholder'ı payload'a sokma

    if not messages:
        sys.exit(0)

    # --- son gerçek user mesajını bul ---
    last_user = None
    for m in reversed(messages):
        if isinstance(m, dict) and m.get("role") == "user":
            last_user = m
            break

    if last_user is None:
        _die(get_ui_text(cfg, "o_No_User_Message"), 1)

    last = last_user
    final_messages = []
    editable_map = {}  # path(str) -> bool

    # Son user mesajını memory sistemine ekle
    last_text = str(last.get("content") or "").strip()

    files_last = last.get("files") or []

    is_pdf_text_only_request = False
    is_pdf_image_only_request = False
    is_pdf_mixed_request = False
    is_pdf_mixed_image_only_request = False

    mixed_pdf_layout = None
    mixed_pdf_source = None

    if isinstance(files_last, list):
        for ff in files_last:
            if not isinstance(ff, dict):
                continue

            p = Path(str(ff.get("path") or ""))
            editable = bool(ff.get("edit"))

            if editable and p.suffix.lower() == ".pdf":
                analysis = _analyze_pdf_kind(p)
                kind = analysis.get("kind")

                if kind == "text_only":
                    is_pdf_text_only_request = True

                elif kind == "mixed":
                    mixed_pdf_source = p
                    mixed_pdf_layout = _extract_pdf_mixed_layout(p)

                    if mixed_pdf_layout and mixed_pdf_layout.get("pages"):
                        is_pdf_mixed_request = True
                    else:
                        # Layout çıkarılamıyorsa güvenli fallback:
                        # metin varsa text-only gibi davran.
                        fallback_text = analysis.get("text") or _safe_read_pdf(p)

                        if len((fallback_text or "").strip()) >= 80:
                            is_pdf_text_only_request = True
                        else:
                            is_pdf_image_only_request = True

                else:
                    is_pdf_image_only_request = True

                break

    if context_mode == "rag":
        # Son user mesajını memory sistemine ekle
        if last_text:
            add_memory_chunk(
                chat_data,
                last_text[:memory_chunk_max_chars],
                source="chat:user"
            )

        # Eğer son mesajda kod dosyası varsa code_context oluştur
        if use_code_context:
            update_code_context_from_message(chat_data, last)

    selected_prompt_blocks = cfg.get("prompt_chooser_blocks", ["copyable"])

    selected_prompt_blocks_set = {
        str(x or "").strip()
        for x in selected_prompt_blocks
        if str(x or "").strip()
    }

    pdf_mode_text = "pdf_text" in selected_prompt_blocks_set
    pdf_mode_image = "pdf_image" in selected_prompt_blocks_set
    pdf_mode_text_image = "pdf_text_image" in selected_prompt_blocks_set

    if is_pdf_mixed_request:
        wants_image = _user_wants_image_edit(last_text)
        wants_text = _user_wants_text_edit(last_text)

        # Kullanıcı text istiyor ama resim aynı kalsın diyorsa:
        # AI'ye image edit yaptırma.
        if wants_text and not wants_image:
            pdf_mode_text = True
            pdf_mode_image = False
            pdf_mode_text_image = False

        is_pdf_mixed_image_only_request = wants_image and not wants_text

    has_editable_pdf_request = (
        is_pdf_text_only_request
        or is_pdf_image_only_request
        or is_pdf_mixed_request
    )

    if has_editable_pdf_request and not (pdf_mode_text or pdf_mode_image or pdf_mode_text_image):
        msg = get_ui_text(cfg, "o_PDF_Mode_Required")

        result_obj = {
            "type": "text",
            "content": msg
        }
        sys.stdout.write(json.dumps(result_obj, ensure_ascii=False))
        sys.stdout.flush()
        sys.exit(0)

    system_prompt = build_global_system_prompt(selected_prompt_blocks, cfg)

    if is_pdf_text_only_request and pdf_mode_text:
        system_prompt += "\n\n" + get_prompt_text(cfg, PDF_RULE_KEYS["pdf_text_only_to_docx"])

    elif is_pdf_image_only_request and pdf_mode_image:
        system_prompt += "\n\n" + get_prompt_text(cfg, PDF_RULE_KEYS["pdf_image"])

    elif is_pdf_mixed_request:
        if pdf_mode_text:
            system_prompt += "\n\n" + get_prompt_text(cfg, PDF_RULE_KEYS["pdf_mixed_text_only"])

        elif pdf_mode_image:
            system_prompt += "\n\n" + get_prompt_text(cfg, PDF_RULE_KEYS["pdf_image"])

        elif pdf_mode_text_image:
            system_prompt += "\n\n" + get_prompt_text(cfg, PDF_RULE_KEYS["pdf_mixed_separate"])

    final_messages.append({
        "role": "system",
        "content": system_prompt
    })

    if context_mode == "rag" and (use_summary or use_retrieval or use_code_context):
        memory_context = build_memory_context(chat_data, last_text)

        if not use_summary:
            # Basit filtre: summary başlığı varsa kaldırmak için memory.py tarafında daha temiz yapılabilir.
            pass

        if memory_context:
            memory_context = memory_context[:summary_max_chars + code_context_max_chars + 4000]

            final_messages.append({
                "role": "system",
                "content": (
                    "Aşağıdaki bilgiler konuşma devamlılığı için verilmiştir. "
                    "Bunlara doğrudan cevap verme, sadece son kullanıcı mesajını anlamak için kullan.\n\n"
                    + memory_context
                )
            })

    cfg = load_config()
    force_lang = bool(cfg.get("force_ui_language", False))

    lang = get_ui_language()
    lang_name = lang_code_to_name(lang)

    if force_lang:
        final_messages.append({
            "role": "system",
            "content": f"Kullanıcı şu dilde cevap vermeni tercih ediyor: {lang_name}"
            })


    if response_style:
        final_messages.append({
            "role": "system",
            "content": f"Kullanıcı şu konuşma tarzını tercih ediyor: {response_style}"
        })

    # ---------------- RECENT MESSAGES MEMORY ----------------
    # Burada sadece önceki son 10 konuşma mesajını ekliyoruz.

    if context_mode == "rag" and use_recent_messages and recent_message_count > 0:
        recent_messages = []

        lookback_count = recent_message_count + 2

        for m in messages[-lookback_count:]:
            if not isinstance(m, dict):
                continue

            if m.get("streaming"):
                continue

            # Son user mesajını tekrar ekleme
            if m is last:
                continue

            role_raw = m.get("role")
            if role_raw not in ("user", "bot", "assistant"):
                continue

            recent_messages.append(m)

        recent_messages = recent_messages[-recent_message_count:]

        if recent_messages:
            final_messages.append({
                "role": "system",
                "content": (
                    "Aşağıdaki mesajlar konuşmanın son kısmıdır. "
                    "Bunlara tek tek cevap verme; sadece en son kullanıcı mesajını anlamak için kullan."
                )
            })

        for m in recent_messages:
            role = _role_map(m.get("role"))

            has_any_attach = False

            if m.get("image") and Path(str(m["image"])).exists():
                has_any_attach = True

            imgs = m.get("images")
            if isinstance(imgs, list) and any(Path(str(p)).exists() for p in imgs):
                has_any_attach = True

            files = m.get("files")
            if isinstance(files, list) and any(
                Path(str(f.get("path") or "")).exists()
                for f in files
                if isinstance(f, dict)
            ):
                has_any_attach = True

            if has_any_attach and include_recent_attachments:
                blocks, _ = _build_blocks_and_cache_info(m, cache_images_dir=None)
                final_messages.append({
                    "role": role,
                    "content": blocks
                })
            else:
                final_messages.append({
                    "role": role,
                    "content": _message_content_with_generated_files(m)
                })
    
    # Cache setup
    REFS_CACHE_DIR.mkdir(parents=True, exist_ok=True)

    # Reference system (TEXT + IMAGE) + CACHE
    if selected_indexes:
        chat_key = _safe_chat_key(chat_file)
        bundle_id = _utc_stamp()
        bundle_dir = REFS_CACHE_DIR / chat_key / bundle_id
        bundle_images_dir = bundle_dir / "images"
        bundle_dir.mkdir(parents=True, exist_ok=True)

        final_messages.append({
            "role": "system",
            "content": (
                "Aşağıdaki mesajlar sadece son kullanıcı mesajını anlaman için referans olarak verilmiştir. "
                "Bu mesajlara cevap verme. Sadece en son kullanıcı mesajını yanıtla."
            )
        })

        cache_items = []
        for i in selected_indexes:
            if i < len(messages) - 1:
                msg = messages[i]
                role = _role_map(msg.get("role"))
                has_any_attach = False

                # editable_map doldur (referans mesajın file edit izinleri)
                files_meta = msg.get("files") or []
                if isinstance(files_meta, list):
                    for ff in files_meta:
                        if not isinstance(ff, dict):
                            continue
                        p_raw = str(ff.get("path") or "").strip()
                        if not p_raw:
                            continue
                        editable_map[p_raw] = bool(ff.get("edit"))

                # images?
                if msg.get("image") and Path(str(msg["image"])).exists():
                    has_any_attach = True
                imgs = msg.get("images")
                if isinstance(imgs, list) and any(Path(str(p)).exists() for p in imgs):
                    has_any_attach = True

                # files?
                files = msg.get("files")
                if isinstance(files, list) and any(Path(str(f.get("path") or "")).exists() for f in files if isinstance(f, dict)):
                    has_any_attach = True

                if has_any_attach:
                    blocks, cached_img = _build_blocks_and_cache_info(msg, bundle_images_dir)
                    final_messages.append({"role": role, "content": blocks})
                else:
                    final_messages.append({"role": role, "content": msg.get("content", "")})
                    cached_img = None

                cache_items.append({
                    "index": i,
                    "role": msg.get("role", ""),
                    "content": msg.get("content", ""),
                    "image_original": msg.get("image", None),
                    "images_original": msg.get("images", None),
                    "files_original": msg.get("files", None),
                    "image_cached": cached_img
                })

        bundle = {
            "bundle_id": bundle_id,
            "created_at_utc": datetime.now(timezone.utc).isoformat(),
            "chat_file": str(chat_file),
            "selected_indexes": selected_indexes,
            "items": cache_items
        }
        try:
            (bundle_dir / "bundle.json").write_text(
                json.dumps(bundle, ensure_ascii=False, indent=2),
                encoding="utf-8"
            )
        except Exception:
            pass

    # Last message (answer) - TEXT + (image / images)
    # last message editable_map doldur
    if isinstance(files_last, list):
        for ff in files_last:
            if not isinstance(ff, dict):
                continue
            p_raw = str(ff.get("path") or "").strip()
            if not p_raw:
                continue
            editable_map[p_raw] = bool(ff.get("edit"))

    # Last message (answer) - TEXT + (images/files)
    role = _role_map(last.get("role"))

    has_any_attach = False

    # images?
    if last.get("image") and Path(str(last["image"])).exists():
        has_any_attach = True

    imgs = last.get("images")
    if isinstance(imgs, list) and any(Path(str(p)).exists() for p in imgs):
        has_any_attach = True

    # files?
    files = last.get("files")
    if isinstance(files, list) and any(
        Path(str(f.get("path") or "")).exists()
        for f in files
        if isinstance(f, dict)
    ):
        has_any_attach = True

    if has_any_attach:
        pdf_mode_for_blocks = "auto"

        if pdf_mode_text:
            pdf_mode_for_blocks = "text"
        elif pdf_mode_image:
            pdf_mode_for_blocks = "image"
        elif pdf_mode_text_image:
            pdf_mode_for_blocks = "auto"

        blocks, _ = _build_blocks_and_cache_info(
            last,
            cache_images_dir=None,
            pdf_mode=pdf_mode_for_blocks
        )
        final_messages.append({"role": role, "content": blocks})
    else:
        final_messages.append({"role": role, "content": last.get("content", "")})

    # Payload
    reply = ""
    usage_obj = None

    try:
        if is_local_model:
            provider, actual_model = _split_local_model(model_name)
            provider_cfg = local_providers.get(provider)

            if not isinstance(provider_cfg, dict):
                _die(get_ui_text(cfg, "o_Local_Provider_Not_Found", provider=provider), 1)

            if not bool(provider_cfg.get("enabled", False)):
                _die(get_ui_text(cfg, "o_Local_Provider_Disabled", provider=provider), 1)

            run_startup = str(provider_cfg.get("run_startup") or "").strip()

            if run_startup:
                try:
                    subprocess.Popen(
                        run_startup,
                        shell=True,
                        stdout=subprocess.DEVNULL,
                        stderr=subprocess.DEVNULL
                    )
                except Exception:
                    pass
            else:
                # chat'e mesaj döndür (stdout)
                result_obj = {
                    "type": "text",
                    "content": get_ui_text(cfg, "o_Start_Local_Server_From_App", provider=provider)
                }
                sys.stdout.write(json.dumps(result_obj, ensure_ascii=False))
                sys.stdout.flush()
                sys.exit(1)

            if provider == "ollama":
                reply = _call_ollama(provider_cfg, actual_model, final_messages)
                parsed_reply = json.loads(reply)
                assistant_text = str(parsed_reply.get("content") or "").strip()

                web_query = _extract_web_search_call(assistant_text)

                if web_query and ("web_search" in selected_prompt_blocks_set):
                    web_results = _tavily_search(cfg, web_query, max_results=5)

                    final_messages.append({
                        "role": "assistant",
                        "content": assistant_text
                    })

                    final_messages.append({
                        "role": "system",
                        "content": (
                            "Web search results are below. "
                            "Answer the user's last question using only these results.\n\n"
                            + web_results
                        )
                    })

                    reply = _call_ollama(provider_cfg, actual_model, final_messages)

                sys.stdout.write(reply)
                sys.stdout.flush()

                try:
                    parsed_reply = json.loads(reply)
                    assistant_text = str(parsed_reply.get("content") or "").strip() if isinstance(parsed_reply, dict) else ""

                    if context_mode == "rag":
                        if assistant_text:
                            add_memory_chunk(
                                chat_data,
                                assistant_text[:memory_chunk_max_chars],
                                source="chat:assistant"
                            )

                        if use_summary and should_update_summary(chat_data, every=summary_update_every):
                           chat_data["summary"] = make_simple_summary(chat_data)[:summary_max_chars]

                        save_chat_data(chat_file, chat_data)

                except Exception:
                    pass

                sys.exit(0)

            else:
                _die(get_ui_text(cfg, "o_Unsupported_Local_Provider", provider=provider), 1)

        # Remote / OpenRouter
        is_image_model = _is_image_generation_model(model_name)

        if (is_pdf_text_only_request or is_pdf_mixed_request) and is_image_model:
            result_obj = {
                "type": "text",
                "content": get_ui_text(cfg, "o_PDF_Has_Text_Select_Text_Model")
            }
            sys.stdout.write(json.dumps(result_obj, ensure_ascii=False))
            sys.stdout.flush()
            sys.exit(0)

        # Text tabanlı PDF düzenleme isteklerinde image model kullanma.
        # Çünkü image model metni düzenlemek yerine görsel üretmeye çalışır.
        try:
            for ff in files_last:
                if not isinstance(ff, dict):
                    continue

                p = Path(str(ff.get("path") or ""))
                if p.suffix.lower() == ".pdf" and bool(ff.get("edit")):
                    has_editable_pdf = True
                    break
        except Exception:
            pass

        selected_prompt_blocks = cfg.get("prompt_chooser_blocks", ["copyable"])
        if not isinstance(selected_prompt_blocks, list):
            selected_prompt_blocks = ["copyable"]

        selected_prompt_blocks = {
            str(x or "").strip()
            for x in selected_prompt_blocks
        }

        has_editable_generated_doc = False
        try:
            for ff in files_last:
                if not isinstance(ff, dict):
                    continue

                p = Path(str(ff.get("path") or ""))
                if p.suffix.lower() in (".docx", ".xlsx", ".txt", ".md", ".pdf") and bool(ff.get("edit")):
                    has_editable_generated_doc = True
                    break
        except Exception:
            pass

        wants_file_create = (
            has_editable_generated_doc
            and ("file_create" in selected_prompt_blocks)
        )

        has_input_images = bool(
            last.get("image")
            or last.get("images")
            or any(
                isinstance(f, dict) and str(f.get("path") or "").lower().endswith((".png", ".jpg", ".jpeg", ".webp"))
                for f in files_last
            )
        )

        should_stream = (
            (not has_input_images)
            and (not wants_file_create)
            and (not is_pdf_text_only_request)
            and (not is_pdf_image_only_request)
            and (not is_pdf_mixed_request)
        )

        web_search_enabled = "web_search" in selected_prompt_blocks

        if web_search_enabled:
            should_stream = False

        payload = {
            "model": model_name,
            "messages": final_messages,
            "max_tokens": 5120,
            "stream": should_stream
        }

        response = requests.post(
            OPENROUTER_URL,
            headers={
                "Authorization": f"Bearer {API_KEY}",
                "Content-Type": "application/json",
                "HTTP-Referer": "http://localhost",
                "X-Title": "capture-ai"
            },
            json=payload,
            timeout=120,
            stream=should_stream
        )

        if response.status_code != 200:
            sys.stderr.write(get_ui_text(cfg, "o_OpenRouter_API_Error", status_code=response.status_code, details=response.text) + "\n")
            sys.exit(1)

        if not should_stream:
            data = response.json()

            usage_data = data.get("usage")
            if usage_data:
                prompt_t = int(usage_data.get("prompt_tokens", 0) or 0)
                comp_t   = int(usage_data.get("completion_tokens", 0) or 0)
                total_t  = int(usage_data.get("total_tokens", 0) or 0)

                usage_obj = {
                    "prompt_tokens": prompt_t,
                    "completion_tokens": comp_t,
                    "total_tokens": total_t
                }

            msg = {}
            try:
                msg = data["choices"][0]["message"]
            except Exception:
                msg = {}

            result_obj = _extract_images_from_message(msg)

            # Bazı image modelleri görseli message içinde değil,
            # response JSON'un başka alanlarında döndürebiliyor.
            # Ama fallback image bulamazsa mevcut text content'i ezmemeli.
            if not result_obj.get("images") and not result_obj.get("images_base64"):
                fallback_images = _extract_images_from_anywhere(data)

                if fallback_images.get("images") or fallback_images.get("images_base64"):
                    old_content = str(result_obj.get("content") or "").strip()

                    result_obj = fallback_images

                    if old_content and not str(result_obj.get("content") or "").strip():
                        result_obj["content"] = old_content

            # AI bazen dosya metadata'sını message.content içinde JSON string olarak döndürür.
            # Onu gerçek result_obj haline çeviriyoruz.
            content_text = str(result_obj.get("content") or "").strip()

            if content_text.startswith("{"):
                try:
                    parsed_content = json.loads(content_text)

                    if isinstance(parsed_content, dict) and (
                        parsed_content.get("generated_files_base64")
                        or parsed_content.get("generated_files_url")
                        or parsed_content.get("generated_files")
                    ):
                        result_obj = parsed_content

                except Exception:
                    pass

            web_query = _extract_web_search_call(content_text)

            if web_query and web_search_enabled:
                if ask_for_web_search and not approved_web_search_query:
                    result_obj = {
                        "type": "text",
                        "content": "",
                        "web_search_request": {
                            "query": web_query,
                            "status": "pending"
                        }
                    }

                    print(json.dumps(result_obj, ensure_ascii=False))
                    return

                query_to_search = approved_web_search_query or web_query
                web_results = _tavily_search(cfg, query_to_search, max_results=5)

                final_messages.append({
                    "role": "assistant",
                    "content": content_text
                })

                final_messages.append({
                    "role": "system",
                    "content": (
                        "Web search results are below. "
                        "Answer the user's last question using only these results. "
                        "If the results are not enough, say that clearly.\n\n"
                        + web_results
                    )
                })

                payload2 = {
                    "model": model_name,
                    "messages": final_messages,
                    "max_tokens": 5120,
                    "stream": False
                }

                response2 = requests.post(
                    OPENROUTER_URL,
                    headers={
                        "Authorization": f"Bearer {API_KEY}",
                        "Content-Type": "application/json",
                        "HTTP-Referer": "http://localhost",
                        "X-Title": "capture-ai"
                    },
                    json=payload2,
                    timeout=120
                )

                if response2.status_code != 200:
                    sys.stderr.write(
                        get_ui_text(
                            cfg,
                            "o_OpenRouter_API_Error",
                            status_code=response2.status_code,
                            details=response2.text
                        ) + "\n"
                    )
                    sys.exit(1)

                data = response2.json()
                msg = data["choices"][0]["message"]
                result_obj = _extract_images_from_message(msg)
                content_text = str(result_obj.get("content") or "").strip()

            # Editable PDF isteğinde AI PDF değil image döndürürse,
            # app image'ı PDF dosyasına paketler.
            generated_files = []

            # 1) PDF text -> AI file_create DOCX verir, app DOCX + PDF üretir
            if is_pdf_text_only_request and pdf_mode_text:
                reply_text = str(result_obj.get("content") or "").strip()
                generated_files = _apply_file_create_from_reply(reply_text)

                clean_text = _remove_named_block(reply_text, "file_create")

                result_obj = {
                    "type": "text",
                    "content": clean_text or get_ui_text(cfg, "o_PDF_Text_Edited"),
                    "generated_files": generated_files,
                }

            # 2) PDF image + scan PDF -> AI PNG verir, app PDF'e çevirir
            elif is_pdf_image_only_request and pdf_mode_image and (
                result_obj.get("images") or result_obj.get("images_base64")
            ):
                image_paths = _save_result_images_to_files(result_obj)

                pdf_out = _image_paths_to_pdf_file(
                    image_paths,
                    output_name="edited-scan-pdf.pdf"
                )

                if pdf_out and pdf_out.exists():
                    generated_files.append({
                        "path": str(pdf_out.resolve()),
                        "name": pdf_out.name
                    })

                result_obj = {
                    "type": "text",
                    "content": get_ui_text(cfg, "o_PDF_Image_Edited"),
                    "generated_files": generated_files,
                }

            # 3) Mixed PDF + sadece image -> AI PNG verir, app eski PDF üstüne image block olarak koyar
            elif is_pdf_mixed_request and pdf_mode_image and (
                result_obj.get("images") or result_obj.get("images_base64")
            ):
                image_paths = _save_result_images_to_files(result_obj)

                pdf_out = None
                if mixed_pdf_source and mixed_pdf_layout:
                    pdf_out = _overlay_images_on_pdf_blocks(
                        mixed_pdf_source,
                        mixed_pdf_layout,
                        image_paths,
                        output_name="edited-pdf-image.pdf"
                    )

                if pdf_out and pdf_out.exists():
                    generated_files.append({
                        "path": str(pdf_out.resolve()),
                        "name": pdf_out.name
                    })

                result_obj = {
                    "type": "text",
                    "content": get_ui_text(cfg, "o_PDF_Image_Block_Edited"),
                    "generated_files": generated_files,
                }

            # 4) Mixed PDF + text only / text image -> AI JSON verir, app layout'u yeniden kurar
            elif is_pdf_mixed_request and (pdf_mode_text or pdf_mode_text_image):
                parsed = None

                content_text = str(result_obj.get("content") or "").strip()

                if content_text.startswith("{"):
                    try:
                        parsed = json.loads(content_text)
                    except Exception:
                        parsed = None

                if isinstance(result_obj, dict) and result_obj.get("type") == "pdf_mixed_rewrite":
                    parsed = result_obj

                if isinstance(parsed, dict) and parsed.get("type") == "pdf_mixed_rewrite":
                    text_replacements = parsed.get("text_replacements") or {}
                    image_replacements = parsed.get("image_replacements") or {}

                    pdf_out = _create_pdf_from_mixed_layout(
                        mixed_pdf_layout or {},
                        text_replacements,
                        image_replacements,
                        output_name="edited-mixed-pdf.pdf"
                    )

                    if pdf_out and pdf_out.exists():
                        generated_files.append({
                            "path": str(pdf_out.resolve()),
                            "name": pdf_out.name
                        })

                    result_obj = {
                        "type": "text",
                        "content": str(parsed.get("content") or get_ui_text(cfg, "o_PDF_Mixed_Edited")),
                        "generated_files": generated_files,
                    }

            if is_pdf_mixed_request and isinstance(result_obj, dict):
                content_text = str(result_obj.get("content") or "").strip()

                if content_text.startswith("{"):
                    try:
                        parsed_content = json.loads(content_text)
                        if isinstance(parsed_content, dict):
                            result_obj = parsed_content
                    except Exception:
                        pass

                if result_obj.get("type") == "pdf_mixed_rewrite":
                    text_replacements = result_obj.get("text_replacements")
                    image_replacements = result_obj.get("image_replacements")

                    # Kullanıcı sadece resim değişikliği istediyse,
                    # AI yazı üretse bile hepsini yok say.
                    if is_pdf_mixed_image_only_request:
                        text_replacements = {}

                    if mixed_pdf_layout and isinstance(text_replacements, dict):
                        out = _create_pdf_from_mixed_layout(
                            mixed_pdf_layout,
                            text_replacements,
                            image_replacements if isinstance(image_replacements, dict) else {},
                            "edited-mixed-pdf.pdf"
                        )

                        if out and out.exists():
                            result_obj = {
                                "type": "text",
                                "content": get_ui_text(cfg, "o_File_Created"),
                                "generated_files": [{
                                    "path": str(out.resolve()),
                                    "name": out.name
                                }]
                            }

                            if usage_obj is not None:
                                result_obj["usage"] = usage_obj

            generated_files = []

            # PDF girdiyse ve AI hâlâ PNG/image döndürdüyse:
            # PNG kullanıcıya gösterilmesin, PDF'e çevrilsin.
            if has_editable_pdf_request and (
                result_obj.get("images")
                or result_obj.get("images_base64")
                or result_obj.get("image")
                or result_obj.get("image_base64")
                or result_obj.get("url")
            ):
                image_paths = _save_result_images_to_files(result_obj)

                pdf_out = None

                if image_paths:
                    if is_pdf_mixed_request and mixed_pdf_source and mixed_pdf_layout:
                        pdf_out = _overlay_images_on_pdf_blocks(
                            mixed_pdf_source,
                            mixed_pdf_layout,
                            image_paths,
                            output_name="edited-pdf-image.pdf"
                        )

                    if not pdf_out:
                        pdf_out = _image_paths_to_pdf_file(
                            image_paths,
                            output_name="edited-pdf-result.pdf"
                        )

                if pdf_out and pdf_out.exists():
                    generated_files.append({
                        "path": str(pdf_out.resolve()),
                        "name": pdf_out.name
                    })

                result_obj = {
                    "type": "text",
                    "content": get_ui_text(cfg, "o_File_Created"),
                    "generated_files": generated_files,
                }

            elif isinstance(result_obj.get("generated_files"), list):
                generated_files = result_obj["generated_files"]

            else:
                generated_files += _apply_ai_returned_files_from_result(result_obj)
                generated_files += _apply_file_create_from_reply(result_obj.get("content", ""))


            if generated_files:
                clean_content = _remove_named_block(
                    result_obj.get("content", ""),
                    "file_create"
                )

                final_text = clean_content.strip()

                if not final_text:
                    final_text = get_ui_text(cfg, "o_File_Created")

                result_obj = {
                    "type": "text",
                    "content": final_text,
                    "generated_files": generated_files
                }

                if usage_obj is not None:
                    result_obj["usage"] = usage_obj

            # hiç image yoksa düz text fallback
            if not result_obj.get("content") and not result_obj.get("images") and not result_obj.get("images_base64"):
                txt = _message_to_plain_text(msg.get("content"))
                result_obj = {
                    "type": "text",
                    "content": fix_mojibake(txt)
                }

            if usage_obj is not None:
                result_obj["usage"] = usage_obj

            reply = json.dumps(result_obj, ensure_ascii=False)

            sys.stdout.write(reply)
            sys.stdout.flush()

        else:
            reply_parts = []
            usage_data = None

            for raw_line in response.iter_lines(decode_unicode=False, chunk_size=1024):
                if not raw_line:
                    continue

                line = raw_line.strip()
                if not line.startswith(b"data:"):
                    continue

                data_str = line[len(b"data:"):].strip()
                if data_str == b"[DONE]":
                    break

                try:
                    evt = json.loads(data_str.decode("utf-8", errors="replace"))
                    if "usage" in evt:
                        usage_data = evt["usage"]
                except Exception:
                    continue

                delta = None
                try:
                    delta = evt["choices"][0]["delta"].get("content")
                except Exception:
                    delta = None

                if delta:
                    delta = fix_mojibake(delta)
                    reply_parts.append(delta)
                    sys.stdout.write(delta)
                    sys.stdout.flush()

            reply = "".join(reply_parts).strip()

            if usage_data:
                prompt_t = int(usage_data.get("prompt_tokens", 0) or 0)
                comp_t   = int(usage_data.get("completion_tokens", 0) or 0)
                total_t  = int(usage_data.get("total_tokens", 0) or 0)

                usage_obj = {
                    "prompt_tokens": prompt_t,
                    "completion_tokens": comp_t,
                    "total_tokens": total_t
                }


            reply = fix_mojibake(reply)
            if not reply:
                sys.stderr.write(get_ui_text(cfg, "o_OpenRouter_Empty_Stream") + "\n")
                sys.exit(1)

    except requests.exceptions.Timeout:
        sys.stderr.write(get_ui_text(cfg, "o_Request_Timeout") + "\n")
        sys.exit(1)

    except requests.exceptions.ConnectionError:
        if is_local_model:
            try:
                provider, _actual_model = _split_local_model(model_name)
            except Exception:
                provider = "local"

            provider_cfg = local_providers.get(provider, {}) if isinstance(local_providers, dict) else {}

            base_url = str(provider_cfg.get("base_url") or "").strip()
            custom_error = str(provider_cfg.get("system_error") or "").strip()

            # custom error varsa onu kullan
            if custom_error:
                msg = custom_error.replace(
                    "{base_url}",
                    base_url or "http://127.0.0.1:11434"
                )
            else:
                # default local error
                msg = get_ui_text(
                    cfg,
                    "o_Local_Connection_Error",
                    provider=provider,
                    base_url=base_url or "http://127.0.0.1:11434"
                )

            # chat formatında döndür
            result_obj = {
                "type": "text",
                "content": msg
            }

            sys.stdout.write(json.dumps(result_obj, ensure_ascii=False))
            sys.stdout.flush()

        else:
            # 🌐 online model
            result_obj = {
                "type": "text",
                "content": get_ui_text(cfg, "o_Internet_Connection_Error")
            }

            sys.stdout.write(json.dumps(result_obj, ensure_ascii=False))
            sys.stdout.flush()

        sys.exit(1)

    except Exception as e:
        sys.stderr.write(get_ui_text(cfg, "o_AI_Error", error=str(e)) + "\n")
        sys.exit(1)

    # ---------------- SAVE ASSISTANT MEMORY ----------------
    if context_mode == "rag":
        try:
            assistant_text = ""

            # image model / local model bazen JSON string döndürüyor
            try:
                parsed_reply = json.loads(reply)
                if isinstance(parsed_reply, dict):
                    assistant_text = str(parsed_reply.get("content") or "").strip()
            except Exception:
                assistant_text = str(reply or "").strip()

            if assistant_text:
                add_memory_chunk(
                    chat_data,
                    assistant_text[:memory_chunk_max_chars],
                    source="chat:assistant"
                )

            if use_summary and should_update_summary(chat_data, every=summary_update_every):
                chat_data["summary"] = make_simple_summary(chat_data)[:summary_max_chars]

            save_chat_data(chat_file, chat_data)

        except Exception:
            pass

    sys.exit(0)

if __name__ == "__main__":
    main()
